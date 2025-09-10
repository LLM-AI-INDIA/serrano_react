import streamlit as st
import os
# os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = r"C:\Dinesh\Projects\Python_projects\video\sarreno_app\sarreno_app\service_account.json"

import pandas as pd
from sqlalchemy import create_engine
import traceback
from google.cloud import bigquery
import pymysql
import shutil
from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv
load_dotenv()
import re



CANON_MAP = {
    # identifiers
    "Medi-Cal ID Number": "Medical ID Number",
    "medical_id_number": "Medical ID Number",
    "youth_name": "Name of the youth",

    # dates / appointments
    "actual_release_date": "Actual release date",
    "scheduled_appointments": "Scheduled Appointments",
    "court_dates": "Court dates",

    # social/economic
    "income_and_benefits": "Income and benefits",
    "food_and_clothing": "Food & Clothing",
    "identification_documents": "Identification documents",
    "life_skills": "Life skills",
    "family_and_children": "Family and children",
    "service_referrals": "Service referrals",
    "home_modifications": "Home Modifications",
    "durable_medical_equipment": "Durable Medical Equipment",
    "Screenings": "Screenings",

    # 🚀 missing fields
    "housing": "Housing",
    "employment": "Employment",
    "transportation": "Transportation",
    "Treatment History": "Treatment History",
    "Treatment History (mental health, physical health, substance use)": "Treatment History",

    # ✅ extra 9 fields from screenshot
    "Race/Ethnicity": "Race/Ethnicity",
    "Residential Address": "Residential Address",
    "Telephone": "Telephone",
    "Medi-Cal health plan assigned": "Medi-Cal health plan assigned",
    "Health Screenings": "Health Screenings",
    "Health Assessments": "Health Assessments",
    "Chronic Conditions": "Chronic Conditions",
    "Prescribed Medications": "Prescribed Medications",
    "Primary physician contacts": "Primary physician contacts",
    "Clinical Assessments": "Clinical Assessments",
    "Emergency contacts": "Emergency contacts"
}

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Rename known variants to canonical column names."""
    if df is None or df.empty:
        return df
    rename_map = {k: v for k, v in CANON_MAP.items() if k in df.columns}
    return df.rename(columns=rename_map)

def normalize_selected_fields(selected_fields):
    """Map UI labels to canonical where needed (e.g., Medi-Cal -> Medical)."""
    return [CANON_MAP.get(f, f) for f in selected_fields]

def get_case_notes(sql_dict, bq_dict, dict_representation):
    """Fetch Case Notes with fallback SQL → BQ → Excel."""
    possible_keys = ["Case Notes", "case_notes", "casenotes"]

    for key in possible_keys:
        if sql_dict.get(key):
            return sql_dict[key]
        if bq_dict.get(key):
            return bq_dict[key]
        if dict_representation.get(key):
            return dict_representation[key]

    return "No case notes available."

def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Rename known variants to canonical column names."""
    if df is None or df.empty:
        return df
    rename_map = {k: v for k, v in CANON_MAP.items() if k in df.columns}
    return df.rename(columns=rename_map)

def normalize_selected_fields(selected_fields):
    """Map UI labels to canonical where needed (e.g., Medi-Cal -> Medical)."""
    return [CANON_MAP.get(f, f) for f in selected_fields]

def get_candidates_by_name(person_input: str):
    """
    Searches Excel, SQL, and BigQuery for all people with the same name.
    Returns a de-duplicated list of formatted strings:
    "Name — Medical ID-XXXX | Telephone Number- XXX | Residential Address- XXX"
    """
    candidates = []

    # Excel
    try:
        file_data = pd.read_excel("ExcelFiles/reentry5.xlsx")
        file_data = normalize_columns(file_data)

        if "Name of the youth" in file_data.columns:
            matches = file_data[
                file_data["Name of the youth"].astype(str).str.strip().str.lower()
                == person_input.strip().lower()
            ]
            for _, row in matches.iterrows():
                name = str(row.get("Name of the youth") or "").strip()
                mid = str(row.get("Medi-Cal ID Number") or "").strip()
                phone = str(row.get("Telephone") or "N/A").strip()
                addr = str(row.get("Residential Address") or "N/A").strip()
                if name and mid:
                    formatted = f"{name} — Medical ID-{mid} | Telephone Number- {phone} | Residential Address- {addr}"
                    candidates.append(formatted)
    except Exception as e:
        print("Excel search error:", e)

    # SQL
    try:
        sql_df = read_cloud_sql(person_input)
        sql_df = normalize_columns(sql_df)
        for _, row in sql_df.iterrows():
            name = str(row.get("youth_name") or "").strip()
            mid = str(row.get("medical_id_number") or "").strip()
            phone = str(row.get("telephone") or "N/A").strip()
            addr = str(row.get("residential_address") or "N/A").strip()
            if name and mid:
                formatted = f"{name} — Medical ID-{mid} | Telephone Number- {phone} | Residential Address- {addr}"
                candidates.append(formatted)
    except Exception as e:
        print("SQL search error:", e)

    # BigQuery
    try:
        bq_df = read_bigquery(person_input)
        bq_df = normalize_columns(bq_df)
        for _, row in bq_df.iterrows():
            name = str(row.get("youth_name") or "").strip()
            mid = str(row.get("medical_id_number") or "").strip()
            phone = str(row.get("telephone") or "N/A").strip()
            addr = str(row.get("residential_address") or "N/A").strip()
            if name and mid:
                formatted = f"{name} — Medical ID-{mid} | Telephone Number- {phone} | Residential Address- {addr}"
                candidates.append(formatted)
    except Exception as e:
        print("BigQuery search error:", e)

    # Deduplicate by Medical ID (Excel → SQL → BigQuery priority)
    unique = {}
    for entry in candidates:
        match = re.search(r"Medical ID-(\d+)", entry)
        if match:
            mid = match.group(1)
            if mid not in unique:
                unique[mid] = entry

    return list(unique.values())


# ✅ BigQuery client
client = bigquery.Client()

# ✅ UI → Actual column names mapping
FIELD_MAP = {
    "Name": "Name of the youth",
    "Medical ID": "Medical ID Number",
    "Release Date": "Actual release date",
    "Appointments": "Scheduled Appointments",
    "Housing": "Housing",
    "Employment": "Employment",
    "Income": "Income and benefits",
    "Food & Clothing": "Food & Clothing",
    "Transportation": "Transportation",
    "ID Docs": "Identification documents",
    "Life Skills": "Life skills",
    "Family": "Family and children",
    "Court Dates": "Court dates",
    "Service Referrals": "Service referrals",
    "Home Modifications": "Home Modifications",
    "Durable Equipment": "Durable Medical Equipment",
    "Case Notes": "Case Notes"
}

def set_table_borders(table, color_rgb=(0, 0, 0)):
    """Apply borders to a table manually (works even without Word styles)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")  # thickness
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "{:02X}{:02X}{:02X}".format(*color_rgb))
        borders.append(border)

    tblPr.append(borders)

# ---------- NEW: Font helpers to force Century Gothic everywhere ----------

def _set_run_font(run, name="Century Gothic", size_pt=None, color_rgb=None):
    """
    Force a run's font (including East Asia paths) to a specific font.
    """
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = rPr.rFonts or OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if rPr.rFonts is None:
        rPr.append(rFonts)

    # python-docx logical name (helps in Word UI)
    run.font.name = name

    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if color_rgb is not None:
        run.font.color.rgb = RGBColor(*color_rgb)


def force_document_font(doc, name="Century Gothic"):
    """
    Apply the desired font to:
      - Normal style (document default)
      - All existing paragraphs/runs
      - All existing tables (headers + cells)
    """
    # Default Normal style
    base = doc.styles['Normal']
    base.font.name = name
    # Ensure East Asia defaults also use the same font
    rPr = base._element.get_or_add_rPr()
    rFonts = rPr.rFonts or OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)
    if rPr.rFonts is None:
        rPr.append(rFonts)

    # Existing paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            _set_run_font(run, name=name)

    # Existing tables (headers + cells)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _set_run_font(run, name=name)

# -------------------------------------------------------------------------


def generate_reentry_care_plan(selected_fields, person_input, app_option):
        

    text = "John Doe — Medical ID-5952280034 | Telephone Number- (555) 271-3237 | Residential Address- 134 Brown St, Los Angeles, CA 90063"

    # Regex patterns
    name_pattern = r"^(.*?)\s+—"
    medical_id_pattern = r"Medical ID-(\d+)"

    # Extract values
    name_match = re.search(name_pattern, text)
    medical_id_match = re.search(medical_id_pattern, text)

    if name_match and medical_id_match:
        name = name_match.group(1)
        medical_id = medical_id_match.group(1)
    try:
        selected_fields = normalize_selected_fields(selected_fields)

        # Excel
        file_data = pd.read_excel("ExcelFiles/reentry5.xlsx")
        file_data = normalize_columns(file_data)

        if medical_id and "Medical ID Number" in file_data.columns:
            person_row = file_data[file_data["Medical ID Number"].astype(str) == str(medical_id)]
        else:
            person_row = file_data[file_data.get("Name of the youth", pd.Series(dtype=str)) == person_input]

        dict_representation = person_row.to_dict(orient="records")[0] if not person_row.empty else {}

        # SQL + BigQuery
        sql_record = read_cloud_sql(person_input, medical_id)
        sql_record = normalize_columns(sql_record)

        bq_record = read_bigquery(person_input, medical_id)
        bq_record = normalize_columns(bq_record)

        # Convert to dicts
        sql_dict = (sql_record.to_dict(orient="records")[0]
                    if isinstance(sql_record, pd.DataFrame) and not sql_record.empty else {})
        bq_dict = (bq_record.to_dict(orient="records")[0]
                   if isinstance(bq_record, pd.DataFrame) and not bq_record.empty else {})

        # Merge dictionaries (Excel → SQL → BQ priority)
        merged_dict = {}
        merged_dict.update(dict_representation)
        merged_dict.update(sql_dict)
        merged_dict.update(bq_dict)
        merged_dict.pop("id", None)

        # ✅ Load Template instead of starting fresh
        doc = Document("D:\Application\data\Template.docx")

        # Title
        doc.add_paragraph("")
        title_text = f"{person_input}'s Reentry Care Plan"
        doc_title = doc.add_paragraph(title_text)
        doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if doc_title.runs:
            run = doc_title.runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        doc.add_paragraph("")

        # Table of all fields
        all_possible_fields = [f for f in dict.fromkeys(CANON_MAP.values()) if f != "Case Notes"]

        table = doc.add_table(rows=1, cols=2)
        

# apply borders directly (XML hack inline)
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')     # thickness
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')  # black
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)


        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"

        for key in all_possible_fields:
            if key in selected_fields:
                value = merged_dict.get(key, "Not Available")
            else:
                value = "Not Selected"

            row_cells = table.add_row().cells
            row_cells[0].text = str(key)
            row_cells[1].text = "" if pd.isna(value) else str(value)

        # Case Notes
        # Case Notes → only if selected
        if "Case Notes" in selected_fields:
            case_notes_value = get_case_notes(sql_dict, bq_dict, dict_representation)
            row_cells = table.add_row().cells
            row_cells[0].text = "Case Notes"
            row_cells[1].text = "" if pd.isna(case_notes_value) else str(case_notes_value).strip()


        # Save as BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        return doc_io

    except Exception as e:
        print("❌ Error in generate_reentry_care_plan:", str(e))
        st.error(f"Failed to generate care plan: {e}")
        return None


def read_cloud_sql(person_input, medical_id=None):
    user = os.environ["CLOUD_SQL_USER"]
    password = os.environ["CLOUD_SQL_PASSWORD"]
    host = os.environ["CLOUD_SQL_HOST"]
    database = "serrano"

    connection_url = f"mysql+pymysql://{user}:{password}@{host}/{database}"
    engine = create_engine(connection_url)

    if medical_id:
        query = f"SELECT * FROM SocialEconomicLogistics_backup WHERE medical_id_number='{medical_id}'"
    else:
        query = f"SELECT * FROM SocialEconomicLogistics_backup WHERE youth_name='{person_input}'"

    df = pd.read_sql(query, engine)
    return df

def read_bigquery(person_input, medical_id=None):
    if medical_id:
        query = """
            SELECT *
            FROM genai-poc-424806.SerranoAdvisorsBQ.scalablefeaturesforBQ
            WHERE medical_id_number = @mid
        """
        job_config = bigquery.QueryJobConfig(
            query_parameters=[bigquery.ScalarQueryParameter("mid", "STRING", str(medical_id))]
        )
    else:
        query = """
            SELECT *
            FROM genai-poc-424806.SerranoAdvisorsBQ.scalablefeaturesforBQ
            WHERE youth_name = @name
        """
        job_config = bigquery.QueryJobConfig(
            query_parameters=[bigquery.ScalarQueryParameter("name", "STRING", person_input)]
        )

    df = client.query(query, job_config=job_config).to_dataframe()
    return df

# ✅ DB CONFIG
DB_CONFIG = {
    "host": "34.44.69.178",
    "user": "root",
    "password": "SQLsql$123",
    "database": "serrano"
}

# ✅ Test DB connection on import
conn = pymysql.connect(**DB_CONFIG)
cursor = conn.cursor()
cursor.execute("SELECT id FROM SocialEconomicLogistics_backup LIMIT 3")
rows = cursor.fetchall()
print("✅ rows:", rows)
cursor.close()
conn.close()