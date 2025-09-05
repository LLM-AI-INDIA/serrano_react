import json
import shutil
from typing import Dict, Any
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- JSON HELPER ----------

def normalize_str_value(value):
    """Convert 'false', 'true', 'null' (as strings) to real Python types."""
    if isinstance(value, str):
        val = value.strip().lower()
        if val == "false":
            return False
        elif val == "true":
            return True
        elif val == "null":
            return None
    return value


def extract_base_info(json_obj: Dict[str, Any]) -> Dict[str, Any]:
    """Dynamically extract candidate info regardless of field naming."""
    base_info = {}
    for key, value in json_obj.items():
        key_lower = key.lower()
        if "candidate" in key_lower:
            base_info["candidate"] = value
        elif "birth" in key_lower:
            base_info["date_of_birth"] = value
        elif "inmate" in key_lower:
            base_info["inmate_number"] = value
    return base_info


# ---------- DOCX HELPERS ----------

def set_table_borders(table):
    """Apply full borders to a Word table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_font(run, font_name="Century Gothic", size=12, bold=False):
    """Set font for a run (paragraph or cell)."""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size)
    run.bold = bold

def add_vertical_table_with_border(doc: Document, section_title: str, data: Dict[str, Any]):
    # Section heading
    heading = doc.add_paragraph()
    run = heading.add_run(section_title)
    set_font(run, font_name="Century Gothic", size=14, bold=True)

    # Table
    table = doc.add_table(rows=0, cols=2)
    for key, value in data.items():
        clean_value = normalize_str_value(value)
        row_cells = table.add_row().cells

        # Apply font to key cell
        p1 = row_cells[0].paragraphs[0]
        run1 = p1.add_run(str(key))
        set_font(run1)

        # Apply font to value cell
        p2 = row_cells[1].paragraphs[0]
        run2 = p2.add_run("" if clean_value is None else str(clean_value))
        set_font(run2)

    set_table_borders(table)
    doc.add_paragraph()


# ---------- MAIN FUNCTION ----------

def json_to_docx_append_vertical_tables(input_json: Dict[str, Any]) -> str:
    """Generate a Word file with vertical tables from input JSON."""
    template_path = "data/Template.docx"
    output_path = "data/output.docx"
    shutil.copy(template_path, output_path)
    doc = Document(output_path)
    
    # Add global heading first
    heading = doc.add_paragraph()
    heading.alignment = 1  # 0=left, 1=center, 2=right, 3=justify
    run = heading.add_run("Health Risk Assessment")
    set_font(run, font_name="Century Gothic", size=18, bold=True)
    run.underline = True
    heading.paragraph_format.space_after = Pt(20)

    # Base candidate info
    base_info = extract_base_info(input_json)
    add_vertical_table_with_border(doc, "Candidate Information", base_info)

    # Iterate over all top-level keys (look for list[dict] to treat as sections)
    for key, value in input_json.items():
        if isinstance(value, list) and all(isinstance(i, dict) for i in value):
            for entry in value:
                merged = {**base_info, **entry}
                add_vertical_table_with_border(doc, key, merged)

    doc.save(output_path)
    print(f"Document saved to {output_path}")
    return output_path